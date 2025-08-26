from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from docx.shared import Inches

vote_show_dic = {
    "отчисления": "За отчисление",
    "замечания":"За замечание",
    "выговора":"За выговор",
    "против":"Против применения дисциплинарного взыскания",
    "воздержались":"Воздержались"
}


def insert_floating_picture(paragraph, image_path, width_in_inches=1, pos_x=2000000, pos_y=1000000):
    from docx.text.paragraph import Paragraph
    import os

    if not isinstance(paragraph, Paragraph):
        raise TypeError(f"Ожидался объект Paragraph, но получен {type(paragraph)}")

    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Файл не найден: {image_path}")

    run = paragraph.add_run()

    # Проверка run.part
    if not hasattr(run, "part"):
        raise AttributeError("run.part отсутствует")

    r_id = run.part.relate_to(image_path, docx.opc.constants.RELATIONSHIP_TYPE.IMAGE, is_external=False)

    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Картинка не найдена: {image_path}")
    run = paragraph.add_run()

    r_id = run.part.relate_to(image_path, docx.opc.constants.RELATIONSHIP_TYPE.IMAGE, is_external=False)

    picxml = f"""
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
                 simplePos="0" relativeHeight="251659264" behindDoc="1"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page">
          <wp:posOffset>{pos_x}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="page">
          <wp:posOffset>{pos_y}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{int(width_in_inches * 914400)}" cy="{int(width_in_inches * 914400 * 0.4)}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="1" name="Picture 1"/>
        <wp:cNvGraphicFramePr>
          <a:graphicFrameLocks noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:nvPicPr>
                <pic:cNvPr id="0" name="Picture 1"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip r:embed="{r_id}"/>
                <a:stretch>
                  <a:fillRect/>
                </a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{int(width_in_inches * 914400)}" cy="{int(width_in_inches * 914400 * 0.4)}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
    """

    drawing = parse_xml(picxml)
    run._r.append(drawing)
def insert_floating_picture(paragraph, image_path, width_in_inches=1, pos_x=2000000, pos_y=1000000):
    from docx.text.paragraph import Paragraph
    import os
    from docx.oxml import parse_xml
    import docx.opc.constants

    if not isinstance(paragraph, Paragraph):
        raise TypeError(f"Ожидался объект Paragraph, но получен {type(paragraph)}")

    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Файл не найден: {image_path}")

    run = paragraph.add_run()

    # Привязываем изображение к текущему run.part
    r_id = run.part.relate_to(
        image_path,
        docx.opc.constants.RELATIONSHIP_TYPE.IMAGE,
        is_external=False
    )

    picxml = f"""
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
                 simplePos="0" relativeHeight="251659264" behindDoc="1"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page">
          <wp:posOffset>{pos_x}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="page">
          <wp:posOffset>{pos_y}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{int(width_in_inches * 914400)}" cy="{int(width_in_inches * 914400 * 0.4)}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="1" name="Picture 1"/>
        <wp:cNvGraphicFramePr>
          <a:graphicFrameLocks noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:nvPicPr>
                <pic:cNvPr id="0" name="Picture 1"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip r:embed="{r_id}"/>
                <a:stretch>
                  <a:fillRect/>
                </a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{int(width_in_inches * 914400)}" cy="{int(width_in_inches * 914400 * 0.4)}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
    """

    drawing = parse_xml(picxml)
    run._r.append(drawing)

def get_show_name(ii):
    if ii in vote_show_dic:
        return vote_show_dic[ii]
    else:
        return ii


def set_font_size(table, font_size):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
def create_hse_docx(year, month, day, date, total_weight, present_weight, number_all, number_now, name, vote_options,
                    vote_result, st0):
    #vote_options = ["Замечание", "Выговор", "против", "отчисление", "воздержаться"]
    #vote_result = ["10 (Акиевававава)", "20 (ываыыуаыуаыу)", "13 (ыуаыуаыуваыуа)", "4 (ыуаыуаыуваыуа)",
    #               "5 (ыуаыуаыуваыуа)"]
    Error = False

    St1 = st0.split("делу")

    if len(St1) > 1:

        narushenue_full = St1[1]
    else:
        narushenue_full = "**НАРУШЕНИЕ**"
        Error = True

    print(narushenue_full)
    #name = narushenue_full.split(",")[0]
    name.strip()
    name = " "+name

    string1 = "Студенческого совета по делу" + narushenue_full

    print("Утверждение позиции " + string1)

    string2 = "Баланова А.С., предложившего рассмотреть дело студент" + narushenue_full

    print("Слушали: " + string2)
    print("Голосовали открытым голососованием:")

    vote_result_num = []
    vote_options_clear = []
    print(vote_result)
    for i in vote_result:
        a = i.split()
        print(i)
        print(a[0])
        vote_result_num.append(int(a[0]))

    print(vote_result_num)

    # for i, j in zip(vote_options, vote_result_num):
    for j in range(len(vote_options)):
        vote_options = list(vote_options)
        i = vote_options[j]
        if i.lower().find("замечан") != -1:
            vote_options_clear.append("замечания")
        elif (i.lower().find("выгово")) != -1:
            vote_options_clear.append("выговора")
        elif (i.lower().find("отчис")) != -1:
            vote_options_clear.append("отчисления")
        elif (i.lower().find("против")) != -1:
            vote_options_clear.append("против")
        elif (i.lower().find("воздерж")) != -1:
            vote_options_clear.append("воздержались")
            vote_result_num[j] = 0
        else:
            vote_options_clear.append(None)
            vote_result_num[j] = 0
    print(vote_options_clear)
    print(f"vote_result_num:{vote_result_num}")

    max_value = max(vote_result_num)
    max_index = vote_result_num.index(max_value)
    result = vote_options_clear[max_index]

    print(f"result: {result}")

    # Создание документа
    document = Document()
    print("Стадия 1")

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.09)  # Верхнее поле
        section.bottom_margin = Cm(0.49)  # Нижнее поле
        section.left_margin = Cm(2.75)  # Левое поле
        section.right_margin = Cm(1.0)

    # Установка шрифта Geramond для всего документа
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'

    print("Стадия 1")



    # Добавляем "шапку" ВШЭ в основной текст

    p1 = document.add_paragraph()
    run1 = p1.add_run(
        "Федеральное\u00A0государственное\u00A0автономное\u00A0образовательное\u00A0учреждение\u00A0образования")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(10)
    run1.italic = True
    run1.bold = True

    p1.paragraph_format.left_indent = Cm(0.03)
    p1.paragraph_format.right_indent = Cm(0.28)
    p1.paragraph_format.space_before = Pt(4.32)
    p1.paragraph_format.space_after = Pt(0)

    # Вторая строка
    p2 = document.add_paragraph()
    run2 = p2.add_run("«Национальный\u00A0исследовательский\u00A0университет\u00A0“Высшая\u00A0школа\u00A0экономики”»")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(10)
    run2.italic = True
    run2.bold = True

    p2.paragraph_format.left_indent = Cm(0.01)
    p2.paragraph_format.right_indent = Cm(0.27)
    p2.paragraph_format.space_before = Pt(0.05)
    p2.paragraph_format.space_after = Pt(0)

    # Третья строка
    p3 = document.add_paragraph()
    run3 = p3.add_run("Студенческий совет НИУ ВШЭ")
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)
    run3.bold = True

    p3.paragraph_format.left_indent = Cm(0)
    p3.paragraph_format.right_indent = Cm(0.27)
    p3.paragraph_format.space_before = Pt(10.85)
    p3.paragraph_format.space_after = Pt(0)

    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    # Один Shift+Enter
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    # Основной текст: Выписка из протокола
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"ВЫПИСКА ИЗ ПРОТОКОЛА № {year}{month:02}{day:02}\nЗАСЕДАНИЯ СТУДЕНЧЕСКОГО CОВЕТА")
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Основной текст: место, дата, и прочее
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Место проведения заседания: ")
    run.font.size = Pt(12)
    run.underline = True
    paragraph.add_run("Интернет").font.size = Pt(12)

    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.right_indent = Cm(9.08)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph()
    run = paragraph.add_run("Дата проведения заседания: ")
    run.font.size = Pt(12)
    run.underline = True
    paragraph.add_run(date).font.size = Pt(12)

    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.right_indent = Cm(9.08)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Число избранных членов Студенческого совета составляет 40 (сорок) человек\n(суммарный вес составляет {total_weight})."
    )

    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.right_indent = Cm(3.03)
    paragraph.paragraph_format.space_before = Pt(13.68)
    paragraph.paragraph_format.space_after = Pt(0)
    run.font.size = Pt(12)


    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"Число присутствующих членов Студенческого совета составляет {number_now} человек (Суммарный вес составляет {present_weight})\n\n"
    )
    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.right_indent = Cm(0.8)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run.font.size = Pt(12)

    if(float(present_weight)>float(total_weight)/2):
        a = document.add_paragraph()
        run = a.add_run("Заседание Студенческого совета правомочно, поскольку совокупный вес голосов присутствующих членов составляет более половины от суммы весов всех членов Студенческого совета.")

        a.paragraph_format.space_after = Pt(8)
        run.italic = True
        a.paragraph_format.left_indent = Cm(0.26)
        a.paragraph_format.right_indent = Cm(0.26)
        a.paragraph_format.space_before = Pt(0)
        a.paragraph_format.space_after = Pt(0)
        a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run.font.size = Pt(12)
    else:
        a = document.add_paragraph()
        run = a.add_run("Заседание Студенческого совета не правомочно, поскольку совокупный вес голосов присутствующих членов составляет менее половины от суммы весов всех членов Студенческого совета.")
        a.paragraph_format.space_after = Pt(8)
        run.italic = True
        a.paragraph_format.left_indent = Cm(0.26)
        a.paragraph_format.right_indent = Cm(0.26)
        a.paragraph_format.space_before = Pt(0)
        a.paragraph_format.space_after = Pt(0)
        a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run.font.size = Pt(12)

    # Повестка дня

    p = document.add_paragraph()
    run = p.add_run("Повестка дня:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(13.75)
    run.font.size = Pt(12)
    run.bold = True
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\tУтверждение позиции Студенческого совета НИУ ВШЭ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run.bold = True
    run.font.size = Pt(12)
    run = paragraph.add_run(f" по делу{narushenue_full}")
    run.font.size = Pt(12)
    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.right_indent = Cm(0.26)


    paragraph = document.add_paragraph()
    run = paragraph.add_run("Голосовали открытым голосованием:")
    run.bold = True
    run.font.size = Pt(12)

    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.right_indent = Cm(0.26)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1.35)
    paragraph.paragraph_format.space_after = Pt(0)


    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.right_indent = Cm(0.26)

    ii = 1

    for option, result_l, clear in zip(vote_options, vote_result, vote_options_clear):

        if clear == None:
            run = paragraph.add_run(f"{option} - ")
            run.bold = True
            run.font.size = Pt(12)
            run = paragraph.add_run(f"{result_l}")
            run.font.size = Pt(12)
            if ii<len(vote_options):
                paragraph.add_run(";").font.size = Pt(12)
            else:
                paragraph.add_run(".").font.size = Pt(12)
            paragraph.add_run(f"\n").font.size = Pt(12)
        else:
            run = paragraph.add_run(f"{get_show_name(clear)} - ")
            run.font.size = Pt(12)
            run.bold = True
            paragraph.add_run(f"{result_l}").font.size = Pt(12)

            if ii<len(vote_options):
                paragraph.add_run(";").font.size = Pt(12)
            else:
                paragraph.add_run(".").font.size = Pt(12)

            paragraph.add_run(f"\n").font.size = Pt(12)
        ii += 1



    #document.add_paragraph()  # Один Shift-Enter
    paragraph = document.add_paragraph()  # Один Enter без Shift
    run = paragraph.add_run("Решили:")
    run.bold = True
    run.font.size = Pt(12)
    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.right_indent = Cm(0.26)



    paragraph = document.add_paragraph()

    paragraph.paragraph_format.left_indent = Cm(0.26)
    paragraph.paragraph_format.right_indent = Cm(0.26)

    run = paragraph.add_run(f"\tРекомендовать")
    run.bold = False
    run.font.size = Pt(12)
    if result == "против":
        run = paragraph.add_run(f" воздержаться от применения мер дисциплинарного взыскания")
        run.bold = True
        run.font.size = Pt(12)
    else:
        run = paragraph.add_run(f" применить")
        run.bold = False
        run.font.size = Pt(12)
    run = paragraph.add_run(f" в отношении{name}")
    run.bold = False
    run.font.size = Pt(12)
    if result != "против":
        run = paragraph.add_run(f" меру дисциплинарного взыскания в виде ")
        run.bold = False
        run.font.size = Pt(12)
        run = paragraph.add_run(str(result) + ".")
        run.bold = True
        run.font.size = Pt(12)


    #if result == "против"
    #paragraph.add_run("замечания.").bold = True
    doc = document
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.26)

    row_data = ["Председательствующий", "__________________________", "Фамилия И.О."]
    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0].cells
    for i, cell_text in enumerate(row_data):
        row[i].text = cell_text


    # Удаление границ таблицы
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')  # Установка правильного пространства имён
                tc_borders.append(border)
            tc_pr.append(tc_borders)
    set_font_size(table, 12)

    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.26)

    row_data = ["Секретарь", "__________________________", "Фамилия И.О."]
    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0].cells
    for i, cell_text in enumerate(row_data):
        row[i].text = cell_text

    # Удаление границ таблицы
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')  # Установка правильного пространства имён
                tc_borders.append(border)
            tc_pr.append(tc_borders)
    set_font_size(table, 12)

    # Сохранение документа
    file_name = f"В2ыписка_из_протокола3_{year}{month:02}{day:02}.docx"
    document.save(file_name)
    print(f"Документ сохранен как {file_name}")
    print("Стадия 1")
    return document


# Пример вызова функции
#create_hse_docx(2025, 1, 27, "15.04.2025", "81919", "60", 40, 22, "___ИМЯ___",
